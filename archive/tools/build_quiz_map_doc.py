import json
import re
from datetime import date
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "tools" / "quiz-map-data.json"
OUTPUT = ROOT / "GrowACrop-Quiz-Map-AI-Handoff.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


class FragmentText(HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts = []
        self.href = None

    def handle_starttag(self, tag, attrs):
        if tag == "a":
            self.href = dict(attrs).get("href")

    def handle_endtag(self, tag):
        if tag == "a" and self.href:
            self.parts.append(f" ({self.href})")
            self.href = None

    def handle_data(self, data):
        self.parts.append(data)

    def text(self):
        return "".join(self.parts).strip()


def html_to_text(fragment):
    parser = FragmentText()
    parser.feed(fragment)
    return parser.text()


def set_font(run, size=11, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, 9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_labeled_paragraph(doc, label, text, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(f"{label}: ")
    set_font(r, 10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text or "Not specified")
    set_font(r, 10.5)
    return p


def add_shopping_item(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{number}. ")
    set_font(r, 9.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_font(r, 9.5)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Title": (28, DARK_BLUE, 0, 8),
        "Subtitle": (13, MUTED, 0, 18),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "GrowACrop | Quiz Map"
    set_font(header.runs[0], 9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_cover(doc, data):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(86)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI HANDOFF REFERENCE")
    set_font(r, 10, bold=True, color=BLUE)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("GrowACrop Quiz Map")

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Complete decision flow, answer keys, recommendation routing, and outcome catalogue")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Canonical source")
    set_font(r, 9, bold=True, color=MUTED)
    p.add_run("\n")
    r = p.add_run(data["source_file"])
    set_font(r, 11, bold=True, color=DARK_BLUE)
    p.add_run("\n")
    r = p.add_run(f"Snapshot prepared {date.today().isoformat()}")
    set_font(r, 9.5, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Use this document as the behavioural specification. Recommendation routing is determined by Step 1 + Step 2.")
    set_font(r, 10.5, italic=True, color=MUTED)
    doc.add_page_break()


def add_overview(doc, data):
    doc.add_heading("1. How the quiz works", level=1)
    add_labeled_paragraph(doc, "Primary routing rule", "recommendation = RECS[selected system][selected Step 2 key]")
    add_labeled_paragraph(doc, "Result-changing answers", "Step 1 (system/style) and Step 2 (system-specific size, budget, plant count, or format).")
    add_labeled_paragraph(doc, "Data-only answer", "Step 3 motivation is collected and sent to the email platform, but it does not change the recommendation.")
    add_labeled_paragraph(doc, "Optional answer", "Step 4 email may be blank. A valid email is submitted to Kit before the result is shown.")
    add_labeled_paragraph(doc, "Result alternatives", "The result page shows every other Step 2 outcome inside the selected Step 1 system. Users can switch among these alternatives without retaking the quiz.")

    doc.add_heading("Decision flow", level=2)
    flow = [
        ("1", "Choose one of 6 setup styles", "Stores A.system"),
        ("2", "Answer the system-specific follow-up", "Stores A.size; this is the second recommendation key"),
        ("3", "Choose one of 5 motivations", "Stores A.motivation; analytics/email field only"),
        ("4", "Optionally enter email", "Stores A.email; validates format; submits to Kit if present"),
        ("Result", "Render RECS[A.system][A.size]", "Show primary recommendation, shopping list, estimated total, and same-system alternatives"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Stage", "User experience", "Implementation meaning"]
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
        shade(table.rows[0].cells[i], LIGHT_BLUE)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            set_font(run, 9.5, bold=True, color=DARK_BLUE)
    set_repeat_header(table.rows[0])
    for row in flow:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for run in cells[i].paragraphs[0].runs:
                set_font(run, 9.5)
    apply_table_geometry(table, [900, 3330, 5130])


def add_questions(doc, data):
    q = data["questions"]
    doc.add_heading("2. Exact question map", level=1)
    doc.add_heading("Step 1 - Choose your style", level=2)
    add_labeled_paragraph(doc, "Prompt", q["step1"]["prompt"])
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, value in enumerate(["Key", "Visible option", "Card crop summary"]):
        table.rows[0].cells[i].text = value
        shade(table.rows[0].cells[i], LIGHT_BLUE)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            set_font(run, 9.5, bold=True, color=DARK_BLUE)
    set_repeat_header(table.rows[0])
    for key, label, grows in q["step1"]["options"]:
        cells = table.add_row().cells
        for i, value in enumerate([key, label, grows]):
            cells[i].text = value
            for run in cells[i].paragraphs[0].runs:
                set_font(run, 9.5)
    apply_table_geometry(table, [1900, 2550, 4910])

    doc.add_heading("Step 2 - Conditional follow-up", level=2)
    rows = []
    prompts = {
        "microgreens": "What space are you working with?",
        "tower": "What space are you working with?",
        "mini_garden": "What's your budget?",
        "dwc_bucket": "How many plants do you want to grow?",
        "salad_table": "Which style suits you best?",
        "grow_shelf": "How much do you want to grow?",
    }
    for system, options in q["step2"].items():
        rows.append((data["system_names"][system], prompts[system], "; ".join(f"{key} = {label}" for key, label in options)))
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, value in enumerate(["Step 1 system", "Step 2 prompt", "Step 2 keys and visible labels"]):
        table.rows[0].cells[i].text = value
        shade(table.rows[0].cells[i], LIGHT_BLUE)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            set_font(run, 9.5, bold=True, color=DARK_BLUE)
    set_repeat_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for run in cells[i].paragraphs[0].runs:
                set_font(run, 9.2)
    apply_table_geometry(table, [1900, 2600, 4860])

    doc.add_heading("Step 3 - Motivation", level=2)
    add_labeled_paragraph(doc, "Prompt", q["step3"]["prompt"])
    for index, (key, label) in enumerate(q["step3"]["options"], 1):
        add_shopping_item(doc, index, f"{label} (key: {key})")
    add_labeled_paragraph(doc, "Routing effect", "None. Motivation is stored and submitted as fields[motivation] when email is provided.")

    doc.add_heading("Step 4 - Optional email", level=2)
    add_labeled_paragraph(doc, "Prompt", q["step4"]["prompt"])
    add_labeled_paragraph(doc, "Helper copy", q["step4"]["note"])
    add_labeled_paragraph(doc, "Validation", "Blank is allowed. Non-blank input must match a basic name@domain.tld pattern.")


def route_label(data, system, size):
    if system == "mini_garden":
        labels = data["labels"]["budget"]
    elif system == "grow_shelf":
        labels = data["labels"]["shelf"]
    elif system == "dwc_bucket":
        labels = data["labels"]["dwc_bucket"]
    elif system == "salad_table":
        labels = data["labels"]["salad_table"]
    else:
        labels = data["labels"]["standard"]
    return labels[size]


def add_routing_matrix(doc, data):
    doc.add_heading("3. Recommendation routing matrix", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["System key", "Step 2 key", "Visible label", "Recommendation", "Headline price"]
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
        shade(table.rows[0].cells[i], LIGHT_BLUE)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            set_font(run, 8.8, bold=True, color=DARK_BLUE)
    set_repeat_header(table.rows[0])
    for system, outcomes in data["recommendations"].items():
        for size, rec in outcomes.items():
            values = [system, size, route_label(data, system, size), rec["name"], rec["price"]]
            cells = table.add_row().cells
            for i, value in enumerate(values):
                cells[i].text = value
                for run in cells[i].paragraphs[0].runs:
                    set_font(run, 8.5)
    apply_table_geometry(table, [1450, 1050, 1650, 3610, 1600])


def split_shop_items(shop):
    return [html_to_text(part) for part in re.split(r"\s+-\s+", shop) if part.strip()]


def add_outcomes(doc, data):
    doc.add_heading("4. Complete outcome catalogue", level=1)
    intro = doc.add_paragraph(
        "Each subsection below is one exact RECS[system][Step 2 key] outcome from the live quiz. "
        "Shopping links are expanded in parentheses so another AI can retain both item copy and destination."
    )
    intro.paragraph_format.space_after = Pt(10)
    outcome_number = 0
    for system, outcomes in data["recommendations"].items():
        doc.add_heading(data["system_names"][system], level=2)
        for size, rec in outcomes.items():
            outcome_number += 1
            doc.add_heading(f"{outcome_number}. {route_label(data, system, size)} -> {rec['name']}", level=3)
            add_labeled_paragraph(doc, "Route", f"{system} -> {size}", after=3)
            add_labeled_paragraph(doc, "Tagline", rec.get("tagline"), after=3)
            add_labeled_paragraph(doc, "Headline price", rec.get("price"), after=3)
            add_labeled_paragraph(doc, "Grows", rec.get("grows"), after=3)
            if rec.get("not_ideal"):
                add_labeled_paragraph(doc, "Not ideal for", rec.get("not_ideal"), after=3)
            add_labeled_paragraph(doc, "Why this recommendation", rec.get("why"), after=5)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run("Shopping list")
            set_font(r, 10.5, bold=True, color=DARK_BLUE)
            for i, item in enumerate(split_shop_items(rec.get("shop", "")), 1):
                add_shopping_item(doc, i, item)
            if outcome_number < 17:
                divider = doc.add_paragraph()
                divider.paragraph_format.space_after = Pt(2)
                p_pr = divider._p.get_or_add_pPr()
                borders = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "4")
                bottom.set(qn("w:space"), "4")
                bottom.set(qn("w:color"), "D9E1EA")
                borders.append(bottom)
                p_pr.append(borders)


def add_behaviour_notes(doc):
    doc.add_heading("5. Result-page and integration behaviour", level=1)
    notes = [
        ("Result title", 'The live UI writes "Your reccomendation" (misspelling preserved in source).'),
        ("Result subtitle", "Based on your choices - [system visible name], [Step 2 visible label]."),
        ("Shopping totals", "Prices are parsed from the shopping-list strings. Optional items are excluded. Grow-light items are excluded from the base total; if a grow-light option exists, both without-LED and with-LED totals are shown."),
        ("Unpriced items", "Items without a parsable dollar amount are excluded from the estimate."),
        ("Quantity handling", 'An item containing "recommend 3" has its first parsed price multiplied by three.'),
        ("Alternative outcomes", "Only other outcomes within the chosen system are shown. Clicking one rerenders it as the primary card."),
        ("Best-match badge", "The badge appears only when the displayed size key equals the original A.size selection. It disappears on an alternative outcome."),
        ("Email endpoint", "POST multipart FormData to https://app.kit.com/forms/9690777/subscriptions."),
        ("Email fields actually sent", "email_address, fields[motivation], and fields[system]. The submit function receives a size label argument but does not append it to FormData."),
        ("Email failure behaviour", "Before-result signup failure blocks result display and asks the user to retry. Result-page email failure leaves the result visible and enables retry."),
        ("Restart", "Clears system, size, motivation, email, selections, result email status, and returns to Step 1."),
    ]
    for label, text in notes:
        add_labeled_paragraph(doc, label, text)

    doc.add_heading("6. Compact machine-readable summary", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Pseudo-logic")
    set_font(r, 10.5, bold=True, color=DARK_BLUE)
    pseudo = [
        "system = Step1.key",
        "variant = Step2[system].key",
        "motivation = Step3.key  // captured, not routed",
        "email = Step4.value     // optional",
        "primary = RECS[system][variant]",
        "alternatives = all RECS[system] entries except variant",
    ]
    for i, line in enumerate(pseudo, 1):
        add_shopping_item(doc, i, line)


def main():
    with DATA.open("r", encoding="utf-8-sig") as f:
        data = json.load(f)
    doc = Document()
    style_document(doc)
    add_cover(doc, data)
    add_overview(doc, data)
    add_questions(doc, data)
    add_routing_matrix(doc, data)
    add_outcomes(doc, data)
    add_behaviour_notes(doc)
    doc.core_properties.title = "GrowACrop Quiz Map - AI Handoff"
    doc.core_properties.subject = "Complete quiz decision map and recommendation catalogue"
    doc.core_properties.author = "GrowACrop"
    doc.core_properties.keywords = "GrowACrop, quiz, decision tree, recommendations, AI handoff"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
