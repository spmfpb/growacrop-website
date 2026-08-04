from docx import Document

path = "GrowACrop-Quiz-Map-AI-Handoff.docx"
doc = Document(path)
texts = [paragraph.text for paragraph in doc.paragraphs]

print("paragraphs", len(texts))
print("tables", len(doc.tables))
print("headings", sum(1 for paragraph in doc.paragraphs if paragraph.style.name.startswith("Heading")))
print("outcome_routes", sum(1 for text in texts if text.startswith("Route: ")))
print("empty_cells", sum(1 for table in doc.tables for row in table.rows for cell in row.cells if not cell.text.strip()))
print("replacement_chars", sum(text.count("\ufffd") for text in texts))
print("first", texts[:8])
print("last", texts[-8:])

assert len(doc.tables) == 4
assert sum(1 for text in texts if text.startswith("Route: ")) == 17
assert not any("\ufffd" in text for text in texts)
assert not any(not cell.text.strip() for table in doc.tables for row in table.rows for cell in row.cells)
print("structural_qa", "PASS")
